
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_to_docx(md_file_path, output_docx_path):
    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_content = []

    for line in lines:
        stripped_line = line.strip()
        
        # Code block handling
        if stripped_line.startswith('```'):
            if in_code_block:
                # End code block
                p = document.add_paragraph()
                run = p.add_run('\n'.join(code_block_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.style = 'No Spacing' # Or create a custom style
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line.rstrip())
            continue

        # Headers
        if stripped_line.startswith('#'):
            level = len(stripped_line.split()[0])
            text = stripped_line.lstrip('#').strip()
            document.add_heading(text, level=min(level, 9))
            continue

        # Horizontal Rule
        if stripped_line.startswith('---') or stripped_line.startswith('***'):
            p = document.add_paragraph()
            run = p.add_run('_' * 30)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue
            
        # Lists (simple)
        if stripped_line.startswith('- ') or stripped_line.startswith('* '):
            text = stripped_line[2:].strip()
            p = document.add_paragraph(text, style='List Bullet')
            apply_bold(p, text)
            continue
            
        if re.match(r'^\d+\.\s', stripped_line):
            text = re.sub(r'^\d+\.\s', '', stripped_line).strip()
            p = document.add_paragraph(text, style='List Number')
            apply_bold(p, text)
            continue
            
        # Blockquotes
        if stripped_line.startswith('> '):
            text = stripped_line[2:].strip()
            p = document.add_paragraph(text)
            p.style = 'Quote' # Or italicize
            continue

        # Normal text (skip empty lines if desired, but keep for spacing)
        if not stripped_line:
            # document.add_paragraph() # Add empty line
            continue

        p = document.add_paragraph()
        apply_bold(p, stripped_line)

    document.save(output_docx_path)
    print(f"Successfully converted {md_file_path} to {output_docx_path}")

def apply_bold(paragraph, text):
    # Simple bold parser for **text**
    # Clear existing runs (hacky, assumes new paragraph)
    paragraph.clear()
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    if len(sys.argv) > 2:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
    else:
        # Default fallback
        md_file = r'd:\360Downloads\zx\接单项目\okx\docs\DEPLOY_GUIDE.md'
        docx_file = r'd:\360Downloads\zx\接单项目\okx\docs\Deployment_Guide.docx'
    
    try:
        md_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"Error: {e}")
